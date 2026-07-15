"""Build deterministic synthetic DOCX and PDF assets for the demo sample."""

from __future__ import annotations

import argparse
import io
import struct
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


REPOSITORY_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_SAMPLE_ROOT = (
    REPOSITORY_ROOT / "samples" / "ja-machine-control-design-review"
)
TARGET_PATH = Path(
    "existing-document/target/basic-design-before-review.docx"
)
REFERENCE_PATH = Path("references/reference-design.docx")
POLICY_PATH = Path("references/quality-assurance-policy.pdf")
OUTPUT_PATHS = (TARGET_PATH, REFERENCE_PATH, POLICY_PATH)

WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
FIXED_CORE_DATETIME = datetime(2026, 1, 1, 0, 0, 0)
BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Noto Sans JP"
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK_BLUE = RGBColor(0x0B, 0x25, 0x45)
MUTED_GRAY = RGBColor(0x5A, 0x64, 0x70)
TABLE_FILL = "F2F4F7"


def _set_font_properties(
    formatting,
    *,
    name: str = BODY_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    formatting.font.name = name
    if size is not None:
        formatting.font.size = Pt(size)
    if color is not None:
        formatting.font.color.rgb = color
    if bold is not None:
        formatting.font.bold = bold

    r_pr = formatting._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute, value in (
        ("w:ascii", name),
        ("w:hAnsi", name),
        ("w:eastAsia", east_asia),
        ("w:cs", name),
    ):
        r_fonts.set(qn(attribute), value)
    for theme_attribute in (
        "w:asciiTheme",
        "w:hAnsiTheme",
        "w:eastAsiaTheme",
        "w:cstheme",
    ):
        r_fonts.attrib.pop(qn(theme_attribute), None)


def _set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    _set_font_properties(
        run,
        size=size,
        color=color,
        bold=bold,
    )


def _set_style_font(
    style,
    *,
    size: float,
    color: RGBColor | None,
    bold: bool | None = None,
) -> None:
    _set_font_properties(
        style,
        size=size,
        color=color,
        bold=bold,
    )


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, size=11, color=None)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 16, 8),
        "Heading 2": (13, HEADING_BLUE, 12, 6),
        "Heading 3": (12, HEADING_DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        _set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = document.styles["Title"]
    _set_style_font(title, size=23, color=RGBColor(0, 0, 0), bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    title_properties = title._element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    subtitle = document.styles["Subtitle"]
    _set_style_font(subtitle, size=13, color=MUTED_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True


def _configure_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    for child in list(numbering):
        numbering.remove(child)

    abstract_number = OxmlElement("w:abstractNum")
    abstract_number.set(qn("w:abstractNumId"), "0")

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), "434C4D31")
    abstract_number.append(nsid)

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract_number.append(multi_level_type)

    template_code = OxmlElement("w:tmpl")
    template_code.set(qn("w:val"), "434C4D31")
    abstract_number.append(template_code)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), BODY_FONT)
    run_fonts.set(qn("w:hAnsi"), BODY_FONT)
    run_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    run_properties.append(run_fonts)
    level.append(run_properties)
    abstract_number.append(level)
    numbering.append(abstract_number)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), "1")
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), "0")
    number.append(abstract_reference)
    numbering.append(number)
    return 1


def _configure_document(document: Document, *, title: str) -> int:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Twips(708)
    section.footer_distance = Twips(708)

    _configure_styles(document)
    bullet_number_id = _configure_numbering(document)

    properties = document.core_properties
    properties.title = title
    properties.subject = "Synthetic CheckListMaker document review demo"
    properties.author = "CheckListMaker Demo"
    properties.last_modified_by = "CheckListMaker Demo"
    properties.category = "Demo"
    properties.comments = "Synthetic, non-sensitive demonstration content"
    properties.keywords = "CheckListMaker, demo, synthetic"
    properties.created = FIXED_CORE_DATETIME
    properties.modified = FIXED_CORE_DATETIME
    properties.revision = 1
    return bullet_number_id


def _add_header_footer(
    document: Document,
    *,
    header_text: str,
    footer_text: str,
) -> None:
    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(header_text)
    _set_run_font(header_run, size=8, color=MUTED_GRAY)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run(footer_text)
    _set_run_font(footer_run, size=8, color=MUTED_GRAY)


def _add_memo_masthead(
    document: Document,
    *,
    title: str,
    subtitle: str,
    metadata: tuple[tuple[str, str], ...],
) -> None:
    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.add_run(title)
    subtitle_paragraph = document.add_paragraph(style="Subtitle")
    subtitle_paragraph.add_run(subtitle)

    # Named memo-masthead override: compact metadata rows at single spacing.
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        label_run = paragraph.add_run(f"{label}: ")
        _set_run_font(label_run, size=10.5, color=RGBColor(0, 0, 0), bold=True)
        value_run = paragraph.add_run(value)
        _set_run_font(value_run, size=10.5, color=RGBColor(0, 0, 0))


def _add_bullet(document: Document, text: str, number_id: int) -> None:
    paragraph = document.add_paragraph()
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_properties.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    number_properties.append(number)
    paragraph_properties.append(number_properties)
    paragraph.add_run(text)


def _set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Twips(width_dxa)
    cell_properties = cell._tc.get_or_add_tcPr()
    width = cell_properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        cell_properties.insert(0, width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    if sum(widths) != 9360:
        raise ValueError("table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.insert(0, table_width)
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = table_properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        table_properties.append(margins)
    for side, value in (
        ("top", 80),
        ("bottom", 80),
        ("start", 120),
        ("end", 120),
    ):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "B8C2CC")

    grid_columns = table._tbl.tblGrid.gridCol_lst
    if len(grid_columns) != len(widths):
        raise ValueError("table grid does not match column count")
    for column, width_dxa in zip(grid_columns, widths, strict=True):
        column.set(qn("w:w"), str(width_dxa))

    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths, strict=True):
            _set_cell_width(cell, width_dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _format_table(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.10
                for run in paragraph.runs:
                    _set_run_font(
                        run,
                        size=10.5,
                        color=INK_BLUE,
                        bold=row_index == 0,
                    )
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), TABLE_FILL)
                cell._tc.get_or_add_tcPr().append(shading)

    row_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    row_properties.append(repeat_header)


def _add_parameter_table(
    document: Document,
    rows: tuple[tuple[str, str], ...],
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "主要パラメータ"
    table.rows[0].cells[1].text = "設計値"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    _set_table_geometry(table, (2700, 6660))
    _format_table(table)


def _patch_utf8_zip_flags(payload: bytes) -> bytes:
    patched = bytearray(payload)
    end_of_central_directory = patched.rfind(b"PK\x05\x06")
    if end_of_central_directory < 0:
        raise ValueError("ZIP end-of-central-directory record missing")
    entry_count = struct.unpack_from(
        "<H", patched, end_of_central_directory + 10
    )[0]
    central_offset = struct.unpack_from(
        "<I", patched, end_of_central_directory + 16
    )[0]

    cursor = central_offset
    for _ in range(entry_count):
        if patched[cursor : cursor + 4] != b"PK\x01\x02":
            raise ValueError("invalid ZIP central-directory entry")
        central_flags = struct.unpack_from("<H", patched, cursor + 8)[0]
        struct.pack_into("<H", patched, cursor + 8, central_flags | 0x0800)
        local_offset = struct.unpack_from("<I", patched, cursor + 42)[0]
        if patched[local_offset : local_offset + 4] != b"PK\x03\x04":
            raise ValueError("invalid ZIP local-file entry")
        local_flags = struct.unpack_from("<H", patched, local_offset + 6)[0]
        struct.pack_into("<H", patched, local_offset + 6, local_flags | 0x0800)

        name_length, extra_length, comment_length = struct.unpack_from(
            "<HHH", patched, cursor + 28
        )
        cursor += 46 + name_length + extra_length + comment_length
    return bytes(patched)


def _normalize_docx(raw_payload: bytes) -> bytes:
    source = io.BytesIO(raw_payload)
    output = io.BytesIO()
    with zipfile.ZipFile(source, "r") as input_archive:
        parts = {
            name: input_archive.read(name)
            for name in input_archive.namelist()
            if not name.endswith("/")
        }
    with zipfile.ZipFile(
        output,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
        strict_timestamps=True,
    ) as output_archive:
        for name in sorted(parts):
            entry = zipfile.ZipInfo(name, FIXED_TIMESTAMP)
            entry.create_system = 3
            entry.external_attr = (0o100644 << 16)
            entry.compress_type = zipfile.ZIP_DEFLATED
            output_archive.writestr(
                entry,
                parts[name],
                compress_type=zipfile.ZIP_DEFLATED,
                compresslevel=9,
            )
    return _patch_utf8_zip_flags(output.getvalue())


def _save_deterministic_docx(document: Document, destination: Path) -> None:
    buffer = io.BytesIO()
    document.save(buffer)
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(_normalize_docx(buffer.getvalue()))


def build_target_docx(destination: Path) -> None:
    destination = Path(destination)
    document = Document()
    bullet_number_id = _configure_document(
        document,
        title="設備状態監視機能 基本設計書",
    )
    _add_header_footer(
        document,
        header_text="BASIC DESIGN REVIEW | INTERNAL",
        footer_text="DMS-2026 | REVIEW BEFORE APPROVAL",
    )
    _add_memo_masthead(
        document,
        title="設備状態監視機能 基本設計書",
        subtitle="レビュー前 | 設計審査用デモ",
        metadata=(
            ("文書管理番号", "DMS-2026"),
            ("版", "0.7"),
            ("機密区分", "社内"),
            ("改訂日", "2026-06-30"),
            ("文書状態", "承認レビュー前"),
        ),
    )

    document.add_heading("1. 目的", level=1)
    document.add_paragraph(
        "設備の温度、振動、および運転状態を周期的に監視し、しきい値を"
        "超えた状態を運転員へ通知する機能を設計する。レビューでは入力取得、"
        "状態判定、通知、および記録の流れが一貫していることを確認する。"
    )

    document.add_heading("2. 適用範囲", level=1)
    document.add_paragraph(
        "対象: 制御盤に接続された温度センサー、振動センサー、および運転状態"
        "信号の取得から、設備状態の判定、警報通知、イベント記録までを含む。"
    )

    document.add_heading("3. 構成", level=1)
    document.add_paragraph(
        "本機能は次の三要素で構成し、状態判定部を中心にデータを受け渡す。"
    )
    _add_bullet(
        document,
        "監視入力部: センサー値と運転状態信号を取得し、時刻情報を付与する。",
        bullet_number_id,
    )
    _add_bullet(
        document,
        "状態判定部: 取得値をしきい値と比較し、正常、警報、異常を判定する。",
        bullet_number_id,
    )
    _add_bullet(
        document,
        "通知出力部: 判定結果を操作画面へ送り、イベント履歴へ記録する。",
        bullet_number_id,
    )

    document.add_heading("4. 機能設計", level=1)
    document.add_paragraph(
        "監視入力部は 500 ms ごとに入力を収集する。状態判定部は現在値と"
        "しきい値を比較し、結果が変化した場合に通知出力部へ判定結果を渡す。"
    )
    document.add_heading("主要パラメータ", level=2)
    _add_parameter_table(
        document,
        (
            ("監視周期", "500 ms"),
            ("警報保持時間", "2 s"),
            ("入力タイムアウト", "1,000 ms"),
        ),
    )

    document.add_heading("5. 異常処理", level=1)
    document.add_paragraph(
        "入力タイムアウトまたはセンサー範囲外を検出した場合は異常として"
        "記録し、運転員へ適切に通知する。入力が正常範囲へ戻った後、連続二回の"
        "正常取得を確認して監視を復旧する。"
    )

    document.add_heading("6. スケジュール", level=1)
    document.add_paragraph(
        "設計レビューを 2026-07-01 に実施し、指摘反映後に改訂版を作成する。"
        "品質保証確認の完了後、承認手続きを開始する。"
    )

    document.add_heading("7. 承認", level=1)
    document.add_paragraph("改訂日: 2026-06-30")
    document.add_paragraph("最終承認者: 未定")
    document.add_paragraph("承認日: 未定")
    document.add_paragraph(
        "承認手順: 設計レビューの指摘を解消し、品質保証担当者の確認後に"
        "最終承認者へ承認を依頼する。"
    )

    _save_deterministic_docx(document, destination)


def build_reference_docx(destination: Path) -> None:
    destination = Path(destination)
    document = Document()
    _configure_document(
        document,
        title="設備状態監視機能 参考設計書",
    )
    _add_header_footer(
        document,
        header_text="REFERENCE DESIGN | SYNTHETIC DEMO",
        footer_text="REF-DESIGN-01 | AUTHORITY: REFERENCE",
    )
    _add_memo_masthead(
        document,
        title="設備状態監視機能 参考設計書",
        subtitle="記述例 | 権威レベル: reference",
        metadata=(
            ("参考資料番号", "REF-DESIGN-01"),
            ("版", "1.0"),
            ("機密区分", "社内"),
            ("作成日", "2026-06-15"),
        ),
    )

    document.add_heading("1. 概要", level=1)
    document.add_paragraph(
        "設備状態監視機能の実装例を示す。本資料は記述例であり、上位の品質"
        "規程または承認済みテンプレートと矛盾する場合は上位資料を優先する。"
    )

    document.add_heading("2. 参考パラメータ", level=1)
    document.add_paragraph(
        "監視入力部は周期的にセンサー値を取得し、状態判定部へ送信する。"
    )
    _add_parameter_table(
        document,
        (
            ("監視周期", "500 ms"),
            ("状態保持時間", "2 s"),
            ("記録方式", "イベント変化時"),
        ),
    )

    document.add_heading("3. 注記", level=1)
    document.add_paragraph(
        "ここに記載した 500 ms は参考値であり、拘束力のある品質規程による"
        "確認を省略してはならない。"
    )

    _save_deterministic_docx(document, destination)


def _draw_pdf_rule(
    pdf: canvas.Canvas,
    *,
    number: str,
    heading: str,
    lines: tuple[str, ...],
    top: float,
) -> float:
    left = 72
    width = letter[0] - 144
    line_height = 17
    height = 35 + line_height * len(lines)
    bottom = top - height
    pdf.setFillColor(HexColor("#F4F6F9"))
    pdf.roundRect(left, bottom, width, height, 5, fill=1, stroke=0)
    pdf.setFillColor(HexColor("#2E74B5"))
    pdf.roundRect(left + 10, top - 28, 24, 20, 4, fill=1, stroke=0)
    pdf.setFillColor(HexColor("#FFFFFF"))
    pdf.setFont("HeiseiKakuGo-W5", 10)
    pdf.drawCentredString(left + 22, top - 22, number)
    pdf.setFillColor(HexColor("#1F4D78"))
    pdf.setFont("HeiseiKakuGo-W5", 12)
    pdf.drawString(left + 44, top - 22, heading)
    pdf.setFillColor(HexColor("#1D2733"))
    pdf.setFont("HeiseiKakuGo-W5", 10.5)
    y = top - 45
    for line in lines:
        pdf.drawString(left + 18, y, line)
        y -= line_height
    return bottom - 10


def build_policy_pdf(destination: Path) -> None:
    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
    pdf = canvas.Canvas(
        str(destination),
        pagesize=letter,
        invariant=1,
        pageCompression=1,
    )
    pdf.setTitle("品質保証規程（デモ）")
    pdf.setAuthor("CheckListMaker Demo")
    pdf.setSubject("基本設計書レビュー用の拘束的な品質規則")
    pdf.setCreator("CheckListMaker Demo")
    pdf.setKeywords("CheckListMaker, synthetic demo, quality assurance")

    page_width, page_height = letter
    pdf.setFillColor(HexColor("#0B2545"))
    pdf.rect(0, page_height - 102, page_width, 102, fill=1, stroke=0)
    pdf.setFillColor(HexColor("#FFFFFF"))
    pdf.setFont("HeiseiKakuGo-W5", 20)
    pdf.drawString(72, page_height - 58, "品質保証規程（デモ）")
    pdf.setFont("HeiseiKakuGo-W5", 9.5)
    pdf.drawString(
        72,
        page_height - 82,
        "基本設計書レビュー用 | 権威レベル: binding | 架空デモ資料",
    )

    pdf.setFillColor(HexColor("#1D2733"))
    pdf.setFont("HeiseiKakuGo-W5", 10)
    pdf.drawString(
        72,
        page_height - 128,
        "本規程は、基本設計書の識別、性能、検証可能性、および承認情報を確認する。",
    )

    top = page_height - 148
    top = _draw_pdf_rule(
        pdf,
        number="01",
        heading="文書識別",
        lines=("管理番号は DMS-[0-9]{4} に一致すること。",),
        top=top,
    )
    top = _draw_pdf_rule(
        pdf,
        number="02",
        heading="性能基準",
        lines=("監視周期は 250 ms 以下とする。",),
        top=top,
    )
    top = _draw_pdf_rule(
        pdf,
        number="03",
        heading="検証可能な表現",
        lines=(
            "「適切に」および「必要に応じて」は検証不能な表現である。",
            "これらの表現は使用を禁止し、判定できる条件または動作を記載する。",
        ),
        top=top,
    )
    top = _draw_pdf_rule(
        pdf,
        number="04",
        heading="機密区分",
        lines=("機密区分は「公開」「社内」「機密」のいずれかとする。",),
        top=top,
    )
    _draw_pdf_rule(
        pdf,
        number="05",
        heading="必須設計情報",
        lines=(
            "設計対象と除外範囲を明示する。",
            "改訂日および最終承認者を明示する。",
        ),
        top=top,
    )

    pdf.setFillColor(HexColor("#5A6470"))
    pdf.setFont("HeiseiKakuGo-W5", 8.5)
    pdf.drawString(72, 38, "管理用デモ資料 | 実在の組織、規程、製品とは関係しない")
    pdf.drawRightString(page_width - 72, 38, "1 / 1")
    pdf.showPage()
    pdf.save()


def build_all(output_root: Path) -> None:
    output_root = Path(output_root)
    build_target_docx(output_root / TARGET_PATH)
    build_reference_docx(output_root / REFERENCE_PATH)
    build_policy_pdf(output_root / POLICY_PATH)


def check_committed(sample_root: Path) -> int:
    sample_root = Path(sample_root)
    with tempfile.TemporaryDirectory() as temporary_directory:
        generated_root = Path(temporary_directory)
        build_all(generated_root)
        mismatches = [
            relative_path
            for relative_path in OUTPUT_PATHS
            if not (sample_root / relative_path).is_file()
            or (sample_root / relative_path).read_bytes()
            != (generated_root / relative_path).read_bytes()
        ]
    for relative_path in mismatches:
        print(relative_path.as_posix())
    return 1 if mismatches else 0


def main(argv) -> int:
    parser = argparse.ArgumentParser(
        description="Build deterministic CheckListMaker demo documents."
    )
    mode = parser.add_mutually_exclusive_group(required=True)
    mode.add_argument("--write", action="store_true")
    mode.add_argument("--check", action="store_true")
    parser.add_argument(
        "--output-root",
        type=Path,
        help="Test-only destination replacing the committed scenario root.",
    )
    arguments = parser.parse_args(argv)
    output_root = arguments.output_root or DEFAULT_SAMPLE_ROOT
    if arguments.write:
        build_all(output_root)
        return 0
    return check_committed(output_root)


if __name__ == "__main__":
    raise SystemExit(main(None))
